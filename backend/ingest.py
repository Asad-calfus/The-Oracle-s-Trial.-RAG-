import logging
import os

from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from backend.config import CHUNK_SIZE, CHUNK_OVERLAP
from backend.ocr import describe_page_image, extract_page_image, is_image_page, page_has_image
from backend.pii import redact_text
from backend.tables import extract_tables_as_markdown
from backend.vectorstore import add_chunks_to_store, delete_document

logger = logging.getLogger(__name__)


# loading the pdf using the pypdfloader lab
def load_pdf(file_path: str, include_images: bool = False):
    """Load a PDF and return one Document per page (text + source/page metadata).

    include_images is an explicit opt-in — vision LLM calls cost money, so
    they only run when the caller (the user, via the upload UI) asks for
    them. When False, pages load exactly as before this feature existed.

    Two cases are handled for image content, both via a vision-LLM
    description:
    - Fully-scanned page (no text layer, has an image) — page_content is
      REPLACED with the description, since there's nothing else there.
    - Mixed page (real text AND a meaningful image/chart) — the description
      is APPENDED to the existing text, so a chart's content becomes
      searchable without discarding the real text already extracted.

    Everything downstream (chunking, embedding, retrieval, citations) then
    works exactly the same, unaware any of the text came from OCR.
    """
    # Normalize to an absolute path so the same file always gets the same
    # "source" metadata, whether it's called with a relative or absolute
    # path — otherwise the same PDF can end up stored twice under two
    # different source values.
    file_path = os.path.abspath(file_path)
    loader = PyPDFLoader(file_path)
    pages = loader.load()
    logger.debug("Loaded %s (%d pages)", file_path, len(pages))

    # Table extraction is pure structural parsing (no LLM call), so unlike
    # image description it's cheap enough to always run, not opt-in.
    for page in pages:
        page_number = page.metadata.get("page", 0)
        tables = extract_tables_as_markdown(file_path, page_number)
        if not tables:
            continue

        logger.info("Page %d of %s has %d table(s)", page_number, file_path, len(tables))
        tables_block = "\n\n".join(tables)
        page.page_content = f"{page.page_content}\n\n[Table content]\n{tables_block}"

    if include_images:
        for page in pages:
            page_number = page.metadata.get("page", 0)
            if not page_has_image(file_path, page_number):
                continue

            scanned = is_image_page(file_path, page_number)
            logger.info(
                "Page %d of %s has image content (scanned=%s), running vision description",
                page_number, file_path, scanned,
            )
            try:
                image_bytes = extract_page_image(file_path, page_number)
                description = describe_page_image(image_bytes)
                if scanned:
                    page.page_content = description
                else:
                    page.page_content = f"{page.page_content}\n\n[Image content]\n{description}"
                logger.info("Vision description succeeded for page %d of %s", page_number, file_path)
            except Exception:
                logger.exception("Vision description failed for page %d of %s", page_number, file_path)

    return pages


def load_text_file(file_path: str) -> list[Document]:
    """Load a .txt or .md file as a single Document.

    No pagination exists for these formats at all, so metadata["page"] is
    None — get_sources() already treats a non-int page as "no page to
    show" (Section 20.2), and the frontend shows just the filename.
    """
    file_path = os.path.abspath(file_path)
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()

    logger.debug("Loaded %s (%d characters, no pagination)", file_path, len(text))
    return [Document(page_content=text, metadata={"source": file_path, "page": None})]


def load_docx_file(file_path: str) -> list[Document]:
    """Load a .docx file as a single Document.

    Word's own "pages" are a rendering-time concern (font, margins, zoom
    all change where a page breaks) — not data stored in the file — so,
    like load_text_file(), this produces one unpaginated Document with
    metadata["page"] = None.
    """
    import docx  # local import: only needed on this path, not at module load

    file_path = os.path.abspath(file_path)
    document = docx.Document(file_path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    logger.debug("Loaded %s (%d paragraphs, no pagination)", file_path, len(document.paragraphs))
    return [Document(page_content=text, metadata={"source": file_path, "page": None})]


def _redact_pages(pages, file_path: str) -> None:
    """Mask structured PII in every page's FINAL text, in place.

    Called once from load_document() for EVERY format — PII redaction is
    plain regex over whatever text a document produced, so it applies
    identically regardless of where that text came from (PDF page, a
    whole .txt file, a .docx's paragraphs). Runs unconditionally: pure
    regex, no LLM cost.
    """
    for page in pages:
        redacted, count = redact_text(page.page_content)
        if count:
            page_number = page.metadata.get("page", 0)
            logger.info("Redacted %d PII match(es) on page %d of %s", count, page_number, file_path)
        page.page_content = redacted


# Extensions this project knows how to read, mapped to nothing here on
# purpose — load_document() below is the single source of truth for which
# extension maps to which loader, so there's only one place to update.
SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".txt", ".md")


def load_document(file_path: str, include_images: bool = False) -> list[Document]:
    """Dispatch to the right loader by file extension, then apply the one
    step every format shares: PII redaction (Section 20.3/20.5).

    Table extraction and image description (OCR) remain PDF-only — both
    are built around pdfplumber's page-level PDF API and don't have an
    equivalent for the other formats in this build.
    """
    file_path = os.path.abspath(file_path)
    extension = os.path.splitext(file_path)[1].lower()

    if extension == ".pdf":
        pages = load_pdf(file_path, include_images=include_images)
    elif extension in (".txt", ".md"):
        pages = load_text_file(file_path)
    elif extension == ".docx":
        pages = load_docx_file(file_path)
    else:
        raise ValueError(f"Unsupported file type: {extension}")

    _redact_pages(pages, file_path)
    return pages


def split_documents(documents):
    """Split page-level Documents into smaller overlapping chunks."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
    )
    chunks = splitter.split_documents(documents)

    # chunk_id gives every chunk a stable identifier, on top of the
    # source/page metadata it already inherited from the page it came from.
    for i, chunk in enumerate(chunks):
        chunk.metadata["chunk_id"] = i

    logger.debug("Split into %d chunks (size=%d, overlap=%d)", len(chunks), CHUNK_SIZE, CHUNK_OVERLAP)
    return chunks


def ingest_document(file_path: str, thread_id: int, include_images: bool = False) -> dict:
    """Run the full ingestion pipeline for one document: load -> chunk -> embed -> store.

    Handles .pdf/.docx/.txt/.md via load_document()'s extension dispatch
    (Section 20) — everything below this point is already format-agnostic,
    since it only ever deals in plain Document objects.

    Every chunk is tagged with thread_id, which is what makes a chat's
    documents searchable only from that same chat — see rag.py's mandatory
    thread filter, which relies on this tag being present.

    Re-uploading a file REPLACES its existing chunks rather than adding a
    second copy beside them — otherwise the same document accumulates
    duplicates every time it's uploaded, inflating both the chunk count and
    the retrieval results.

    Returns {"chunks": int, "replaced": int} so the /upload endpoint can tell
    the user what actually happened.
    """
    file_path = os.path.abspath(file_path)
    logger.info(
        "Ingesting %s for thread_id=%s (include_images=%s)",
        file_path, thread_id, include_images,
    )

    replaced = delete_document(file_path)
    if replaced:
        logger.debug("Replaced %d existing chunks before re-ingesting", replaced)

    documents = load_document(file_path, include_images=include_images)
    chunks = split_documents(documents)

    for chunk in chunks:
        chunk.metadata["thread_id"] = thread_id

    add_chunks_to_store(chunks)
    logger.info(
        "Ingested %s: %d chunks tagged thread_id=%s (%d replaced)",
        os.path.basename(file_path), len(chunks), thread_id, replaced,
    )

    return {"chunks": len(chunks), "replaced": replaced}
